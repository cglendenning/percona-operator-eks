#!/usr/bin/env python3
"""
Generates the On-Premises Data Platform Disaster Recovery Word document.
Output: docs/on-prem-dr-narrative.docx

Run: python3 docs/generate_dr_word_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = os.path.join(os.path.dirname(__file__), "on-prem-dr-narrative.docx")

# ── Colour palette ───────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x36, 0x64)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_GREY = RGBColor(0xD9, 0xD9, 0xD9)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
RED_WARN   = RGBColor(0xC0, 0x00, 0x00)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    return p


def add_para(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(11)
        p.add_run(" " + text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_note(doc, text):
    """Italicised note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run("Note: " + text)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    p.paragraph_format.space_after = Pt(6)
    return p


def header_row(table, *cols, bg=DARK_BLUE, fg=WHITE):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = text
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = fg
                run.font.size = Pt(10)


def add_table_row(table, *values, shade=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        if shade:
            set_cell_bg(cell, LIGHT_GREY)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    return row


def add_page_break(doc):
    doc.add_page_break()


# ── Document factory ─────────────────────────────────────────────────────────

def build_document() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default body font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("[ORGANIZATION_NAME]")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()  # spacer

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("On-Premises Data Platform\nDisaster Recovery —\nNarrative and Architecture Commentary")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    meta_tbl = doc.add_table(rows=9, cols=2)
    meta_tbl.style = "Table Grid"
    meta_data = [
        ("Document Title",       "On-Premises Data Platform Disaster Recovery — Narrative and Architecture Commentary"),
        ("Document ID",          "DR-PLN-002"),
        ("Version",              "1.0"),
        ("Classification",       "Internal — Restricted"),
        ("Owner",                "Infrastructure & Platform Engineering"),
        ("Prepared By",          "[AUTHOR_NAME]"),
        ("Review Cycle",         "Annually or after any declared disaster event"),
        ("Applicable Frameworks","CFIUS National Security Review; CA/Browser Forum Certificate Management"),
        ("Last Reviewed",        "2026-06-17"),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_tbl.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_bg(row.cells[0], MID_BLUE)
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    # Relationship note
    rel_p = doc.add_paragraph()
    r = rel_p.add_run(
        "Relationship to DR-PLN-001: "
    )
    r.font.bold = True
    r.font.size = Pt(10)
    r = rel_p.add_run(
        "This document (DR-PLN-002) is a standalone on-premises narrative companion to the "
        "Platform Engineering operational runbooks hosted in the DR Emergency Response Dashboard "
        "(see Section 9). The Dashboard at [DR_DASHBOARD_URL] is the authoritative source for "
        "step-by-step recovery procedures. This document provides architectural commentary, "
        "SeaweedFS backup infrastructure detail, compliance narrative, and auditor-facing "
        "context that is auxiliary to — and explicitly does not supersede — those runbooks."
    )
    r.font.size = Pt(10)
    r.font.italic = True

    add_page_break(doc)

    # ── APPROVALS ───────────────────────────────────────────────────────────
    add_heading(doc, "Approvals", level=1)
    appr_tbl = doc.add_table(rows=5, cols=4)
    appr_tbl.style = "Table Grid"
    header_row(appr_tbl, "Role", "Name", "Signature", "Date")
    for role in [
        "Platform Engineering Lead",
        "Chief Information Security Officer",
        "Change Advisory Board Chair",
        "Data Protection Officer",
    ]:
        add_table_row(appr_tbl, role, "", "", "")
    doc.add_paragraph()
    add_page_break(doc)

    # ── SECTION 1: EXECUTIVE SUMMARY ─────────────────────────────────────
    add_heading(doc, "1  Executive Summary", level=1)
    add_para(doc,
        "This document provides the authoritative on-premises narrative, architectural "
        "commentary, and compliance evidence for the data platform's Disaster Recovery "
        "programme. The platform delivers managed relational database services on "
        "on-premises VMware-hosted Kubernetes infrastructure, powered by Percona XtraDB "
        "Cluster (PXC) under the Percona Kubernetes Operator, with SeaweedFS-based "
        "distributed object storage providing S3-compatible backup retention and "
        "cross-site data protection."
    )
    add_para(doc,
        "This document has been prepared specifically for the on-premises operating "
        "environment. It does not describe or reference any public cloud infrastructure. "
        "All storage, compute, and network components described herein reside within "
        "data centres under the direct physical and administrative control of "
        "[ORGANIZATION_NAME]."
    )
    add_para(doc,
        "The step-by-step recovery runbooks for each disaster scenario are maintained "
        "exclusively in the DR Emergency Response Dashboard (Section 9). This document "
        "references those scenarios by name and provides architectural context, risk "
        "commentary, and compliance observations that auditors and governance bodies "
        "require but that do not belong in an operational runbook."
    )
    add_para(doc, "Key risk posture summary:")
    bullets_exec = [
        ("Zero data loss", "is achievable for the majority of failure scenarios due to "
         "synchronous Galera replication within the PXC cluster."),
        ("Point-in-time recovery (PITR)", "to within 60 seconds of any transaction is "
         "operational via continuous binary log uploads to SeaweedFS-backed S3 storage."),
        ("Full data-centre loss", "is recoverable within 30 minutes with up to 60 seconds "
         "of data loss, contingent on the health of the asynchronous DR replica at dr-std."),
        ("SeaweedFS resilience", "is provided by a 3-master / 3-volume-server topology "
         "with a dedicated Filer instance per cluster and asynchronous replication "
         "controllers at both the standard (dr-std) and secondary (dr-sec) DR sites."),
        ("Critical scenarios", "(ransomware, accidental restore) have RTOs between "
         "4–8 hours with documented and reviewed runbooks."),
        ("Certificate failures", "carry a 45-minute RTO with a tested rotation procedure "
         "governed under the CA/Browser Forum (CAB) Certificate Management process."),
    ]
    for bold_part, rest in bullets_exec:
        add_bullet(doc, rest, bold_prefix=bold_part)

    add_page_break(doc)

    # ── SECTION 2: PURPOSE, SCOPE, APPLICABILITY ────────────────────────────
    add_heading(doc, "2  Purpose, Scope, and Applicability", level=1)
    add_heading(doc, "2.1  Purpose", level=2)
    add_para(doc,
        "This document establishes the formal on-premises architectural narrative, "
        "compliance evidence base, and scenario reference catalogue for the data "
        "platform Disaster Recovery programme. It defines recovery objectives, "
        "documents the SeaweedFS storage architecture, assigns responsibilities, "
        "and provides the audit evidence that regulators and governance bodies require "
        "to assess operational resilience — while explicitly deferring all procedural "
        "detail to the DR Emergency Response Dashboard."
    )
    add_heading(doc, "2.2  Scope", level=2)
    add_para(doc, "This plan covers the following on-premises components:")
    scope_items = [
        "Percona XtraDB Cluster (PXC) instances in the primary data centre and both DR sites.",
        "HAProxy and ProxySQL load-balancing and connection-routing tiers.",
        "SeaweedFS distributed object storage (Masters, Volume Servers, Filer) providing "
        "S3-compatible backup targets and POSIX file semantics.",
        "SeaweedFS asynchronous replication controllers at the dr-std and dr-sec DR sites.",
        "Kubernetes control planes (VMware-hosted, managed by Rancher).",
        "Percona Kubernetes Operator managing the lifecycle of all PXC cluster resources.",
        "Istio service mesh providing mutual TLS (mTLS) between all cluster services.",
        "Percona Monitoring and Management (PMM) providing observability and alerting.",
        "The DR Emergency Response Dashboard web application.",
        "Fleet/Rancher GitOps pipeline used for declarative configuration management.",
    ]
    for item in scope_items:
        add_bullet(doc, item)
    add_para(doc,
        "This document is scoped exclusively to the on-premises operating environment. "
        "No public cloud infrastructure is referenced or implied. Application-layer "
        "software, network infrastructure outside the Kubernetes cluster boundary, and "
        "general IT disaster recovery for non-data-platform systems are out of scope."
    )
    add_heading(doc, "2.3  Applicability", level=2)
    add_para(doc,
        "This document is applicable to: all engineering personnel with operational "
        "responsibility for the data platform; the Security Operations team and any "
        "third-party managed security service provider with access to platform "
        "infrastructure; the Change Advisory Board (CAB) with respect to certificate "
        "management changes; and external auditors, regulators, and reviewers operating "
        "under CFIUS, SOC 2, ISO 27001, or equivalent frameworks."
    )
    add_page_break(doc)

    # ── SECTION 3: DEFINITIONS ───────────────────────────────────────────────
    add_heading(doc, "3  Definitions and Abbreviations", level=1)
    defs_tbl = doc.add_table(rows=1, cols=2)
    defs_tbl.style = "Table Grid"
    header_row(defs_tbl, "Term", "Definition")
    defs = [
        ("CAB", "CA/Browser Forum Certificate Management. The governance body that reviews and "
                "approves changes to certificate infrastructure, including root CA rotations."),
        ("CFIUS", "Committee on Foreign Investment in the United States. A federal interagency "
                  "committee that reviews certain foreign investments for national security "
                  "implications."),
        ("DR", "Disaster Recovery. The set of policies, tools, and procedures to enable the "
               "recovery of technology infrastructure after a disaster."),
        ("DRP", "Disaster Recovery Plan. This document (DR-PLN-002) plus the associated "
                "Dashboard runbooks."),
        ("Filer", "The SeaweedFS Filer service — a stateful component that provides POSIX-like "
                  "file-system metadata semantics and an S3-compatible API gateway on top of "
                  "SeaweedFS Volume Servers."),
        ("Fleet", "Rancher Fleet — a GitOps continuous-delivery tool used for declarative "
                  "Kubernetes configuration management across all on-premises clusters."),
        ("Galera", "The synchronous multi-primary replication library used by PXC."),
        ("GitOps", "An operational model using Git as the single source of truth for declarative "
                   "infrastructure configuration, with automated reconciliation."),
        ("HAProxy", "High Availability Proxy. TCP/HTTP load balancer routing client connections "
                    "to PXC."),
        ("IST", "Incremental State Transfer. The preferred PXC method for re-syncing a node "
                "with a short gap, using the joiner's local data."),
        ("Istio", "An open-source service mesh providing mutual TLS, traffic management, and "
                  "observability between cluster services."),
        ("MTTR", "Mean Time To Recovery. The expected elapsed time from incident detection to "
                 "full service restoration."),
        ("mTLS", "Mutual TLS. Both client and server authenticate each other using X.509 "
                 "certificates."),
        ("Nix", "A purely functional package manager and build system used for declarative, "
                "reproducible infrastructure configuration."),
        ("PITR", "Point-in-Time Recovery. The ability to restore a database to any arbitrary "
                 "moment in the past by replaying binary logs on top of a backup."),
        ("PMM", "Percona Monitoring and Management. The observability platform for database and "
                "Kubernetes metrics."),
        ("ProxySQL", "An advanced MySQL proxy providing connection multiplexing and query "
                     "routing."),
        ("PVC", "PersistentVolumeClaim. A Kubernetes resource representing a request for "
                "durable storage."),
        ("PXC", "Percona XtraDB Cluster. A fully open-source, enterprise-class MySQL cluster "
                "solution."),
        ("Rancher", "A multi-cluster Kubernetes management platform used to manage all "
                    "on-premises Kubernetes clusters."),
        ("RPO", "Recovery Point Objective. Maximum tolerable data loss measured in time."),
        ("RTO", "Recovery Time Objective. Maximum tolerable duration of a service outage."),
        ("SeaweedFS", "A distributed file system and object store providing POSIX and S3 APIs, "
                      "used for blob storage and backup retention in the on-premises environment."),
        ("SIEM", "Security Information and Event Management. Aggregates and correlates log data "
                 "for security monitoring."),
        ("SST", "State Snapshot Transfer. PXC fallback for re-syncing a node when IST is not "
                "possible; transfers the full dataset from a donor node."),
        ("SVID", "SPIFFE Verifiable Identity Document. Short-lived X.509 certificates issued by "
                 "Istio's CA to workloads."),
        ("VMware", "The hypervisor platform hosting all Kubernetes nodes in the on-premises "
                   "environment."),
        ("wsrep", "Write-Set Replication. The Galera protocol underlying PXC synchronous "
                  "replication."),
    ]
    for i, (term, definition) in enumerate(defs):
        add_table_row(defs_tbl, term, definition, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_page_break(doc)

    # ── SECTION 4: ARCHITECTURE ──────────────────────────────────────────────
    add_heading(doc, "4  On-Premises System Architecture", level=1)
    add_heading(doc, "4.1  Platform Overview", level=2)
    add_para(doc,
        "The data platform provides relational database-as-a-service to consuming "
        "applications across an on-premises multi-data-centre topology. The platform "
        "runs Percona XtraDB Cluster (PXC) topologies managed by the Percona Kubernetes "
        "Operator, hosted on VMware-backed Kubernetes clusters managed by Rancher. All "
        "declarative configuration is maintained in version-controlled Git and applied "
        "via the Rancher Fleet GitOps pipeline. All changes to infrastructure, including "
        "certificate rotations, pass through this pipeline and are subject to the "
        "organisation's change management process."
    )
    add_para(doc,
        "Persistent backup storage is provided by SeaweedFS — a distributed object "
        "store deployed on-premises that exposes an S3-compatible API. SeaweedFS "
        "replaces dependence on public cloud object storage and ensures that all backup "
        "artefacts remain within organisation-controlled infrastructure."
    )

    add_heading(doc, "4.2  Physical and Logical Topology", level=2)
    add_para(doc,
        "The on-premises deployment spans three sites:"
    )
    sites = [
        ("Primary Data Centre", "Hosts the active PXC cluster, SeaweedFS primary namespace, "
         "Istio control plane, HAProxy, ProxySQL, and the DR Dashboard. All production "
         "write traffic is served from this site."),
        ("Standard DR Site (dr-std)", "Hosts an asynchronous PXC replica that receives "
         "continuous binary log replication from the primary cluster. Also hosts a "
         "SeaweedFS async replication controller pod that maintains object storage "
         "consistency from the primary site's Filer. This is the primary failover target "
         "for a declared disaster event."),
        ("Secondary DR Site (dr-sec)", "Hosts a second asynchronous PXC replica and a "
         "second SeaweedFS async replication controller pod. Provides geographic and "
         "regulatory separation from dr-std and serves as the tertiary failover option. "
         "Designed to satisfy CFIUS requirements for data sovereignty within "
         "organisation-controlled premises."),
    ]
    for name, desc in sites:
        add_bullet(doc, desc, bold_prefix=name)
    add_para(doc,
        "The primary cluster PXC topology consists of a three-node synchronous Galera "
        "cluster, ensuring that no transaction is acknowledged to clients until it has "
        "been applied to a quorum of nodes. HAProxy provides both a read-only "
        "round-robin endpoint and a single logical write endpoint. PodDisruptionBudgets "
        "and Kubernetes topology spread constraints prevent co-scheduling of PXC pods "
        "on the same VMware host, eliminating single-node failure as a write-path risk."
    )

    add_heading(doc, "4.3  SeaweedFS Distributed Storage Architecture", level=2)
    add_para(doc,
        "SeaweedFS is the on-premises distributed object storage platform that provides "
        "the S3-compatible backup target for all PXC backup operations. It replaces "
        "dependence on public cloud object storage and is the primary resilience "
        "component for backup data sovereignty. This section provides a detailed "
        "architectural description of the SeaweedFS deployment."
    )

    add_heading(doc, "4.3.1  Component Architecture", level=3)
    add_para(doc,
        "The production SeaweedFS deployment uses the following component topology per "
        "primary-site cluster namespace:"
    )
    comp_tbl = doc.add_table(rows=1, cols=4)
    comp_tbl.style = "Table Grid"
    header_row(comp_tbl, "Component", "Count", "Port(s)", "Role")
    components = [
        ("Master Server", "3",
         "9333 (HTTP API)\n19333 (gRPC)",
         "Coordinates volume allocation, manages cluster topology, provides HA "
         "consensus via Raft. Three Masters form a quorum — one acts as leader "
         "at any time. The leader accepts volume assignment requests from Volume "
         "Servers and Filer."),
        ("Volume Server", "3",
         "8080 (HTTP API)\n18080 (gRPC)",
         "Stores actual data blobs. Each file uploaded to SeaweedFS is split into "
         "one or more needles and written to Volume Servers. The replication "
         "factor is configured at the rack/data-centre level via the Master's "
         "replication policy (production minimum: 2 replicas across 2 Volume "
         "Servers). Provides raw byte storage with O(1) disk seeks via direct "
         "needle addressing."),
        ("Filer", "1 per cluster",
         "8888 (HTTP API/POSIX)\n18888 (gRPC)\n8333 (S3 API)",
         "Provides POSIX-like file-system semantics and metadata routing on top of "
         "the Volume Server tier. The Filer maintains a metadata store (filer.toml) "
         "that maps file paths to volume needle IDs. It also exposes an "
         "S3-compatible REST API on port 8333. This S3 endpoint is the target for "
         "all Percona XtraDB Cluster backup jobs (PBM/xtrabackup). The Filer "
         "stores filer metadata in a persistent, crash-recoverable store (LevelDB "
         "or filer-embedded store per deployment configuration)."),
    ]
    for row_data in components:
        add_table_row(comp_tbl, *row_data)
    doc.add_paragraph()

    add_heading(doc, "4.3.2  SeaweedFS as the Backup Target", level=3)
    add_para(doc,
        "Percona Backup for MySQL (PBM), orchestrated by the Percona Kubernetes "
        "Operator, writes all backup artefacts to the SeaweedFS Filer S3 endpoint. "
        "The backup storage configuration references the Filer's Kubernetes service "
        "DNS name:"
    )
    add_para(doc,
        "    S3 endpoint: http://seaweedfs-filer.[seaweedfs-namespace].svc.cluster.local:8333",
        bold=False, italic=False, size=10
    )
    add_para(doc,
        "The backup bucket (e.g., percona-backups) is created in the SeaweedFS Filer "
        "and behaves identically to an S3 bucket from the perspective of the Percona "
        "Operator. Backup credentials (access key and secret key) are stored as "
        "Kubernetes Secrets in the percona namespace and rotated via the Fleet GitOps "
        "pipeline."
    )
    add_para(doc,
        "SeaweedFS provides two properties critical to backup integrity:"
    )
    add_bullet(doc,
        "Append-only needle writes: Volume Server needles are append-only by design. "
        "Once a needle is written, its offset in the volume file is immutable. This "
        "provides a natural write-once property for backup objects, which is a "
        "prerequisite for ransomware resistance.",
        bold_prefix="Write-once semantics:"
    )
    add_bullet(doc,
        "Each volume can be configured with a replication factor (e.g., rack:001 — "
        "one copy on each of two racks). In the production 3-volume-server topology, "
        "backup data is replicated across at least two Volume Servers before the Filer "
        "acknowledges the write to the backup client. This prevents a single volume "
        "server failure from creating a backup gap.",
        bold_prefix="Intra-site replication:"
    )

    add_heading(doc, "4.3.3  Cross-Site Replication Architecture", level=3)
    add_para(doc,
        "SeaweedFS backup artefacts and file namespace state are replicated "
        "asynchronously from the primary site to both DR sites. This is accomplished "
        "via dedicated replication controller pods co-deployed with the DR-site "
        "Kubernetes workloads:"
    )
    repl_tbl = doc.add_table(rows=1, cols=3)
    repl_tbl.style = "Table Grid"
    header_row(repl_tbl, "Site", "Component", "Function")
    repl_rows = [
        ("Primary DC", "SeaweedFS Filer (seaweedfs-primary namespace)",
         "Origin for all backup writes. Exposes S3 API for PXC backup jobs and "
         "Filer HTTP API for replication clients."),
        ("dr-std", "SeaweedFS Async Replication Controller Pod",
         "Continuously monitors the primary Filer for new or updated objects "
         "and replicates them to the dr-std Filer instance. Maintains replication "
         "state and provides recovery hooks for failover scenarios."),
        ("dr-sec", "SeaweedFS Async Replication Controller Pod",
         "Independent replication controller targeting the dr-sec Filer instance. "
         "Provides geographic and regulatory separation independent of dr-std."),
    ]
    for i, row_data in enumerate(repl_rows):
        add_table_row(repl_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_para(doc,
        "The replication is asynchronous. In normal operation, replication lag is "
        "bounded by network bandwidth and write volume. Replication lag is monitored "
        "by PMM; sustained lag beyond a configurable threshold triggers an alert. "
        "During a declared primary-DC disaster event, the DR site's replicated "
        "SeaweedFS Filer is promoted to primary and the PXC backup jobs are "
        "re-pointed to the DR-site Filer endpoint."
    )
    add_note(doc,
        "Cross-site replication does not provide synchronous consistency. A very "
        "recent backup that was written to the primary Filer but not yet replicated "
        "to dr-std at the moment of primary DC loss may need to be recovered via "
        "PITR from the binlog stream if it is within the retention window."
    )

    add_heading(doc, "4.3.4  Credential and Access Control Model", level=3)
    add_para(doc,
        "SeaweedFS S3 credentials are provisioned as follows:"
    )
    cred_items = [
        "Credentials are defined in the SeaweedFS Filer configuration (s3.config or "
        "via weed shell s3.configure) and stored in a Kubernetes Secret in the "
        "seaweedfs namespace.",
        "The Percona Operator backup Secret (percona-backup-seaweedfs-credentials or "
        "equivalent) in the percona namespace holds the access key and secret key used "
        "by PBM/xtrabackup to write to the Filer S3 endpoint.",
        "Credentials follow the principle of least privilege: backup service accounts "
        "have PUT/GET/LIST permissions on the backup bucket only.",
        "Credential rotation is executed via the Fleet GitOps pipeline and requires a "
        "pull-request review. Emergency rotation bypasses the review gate but requires "
        "a post-hoc audit trail entry.",
        "All credential rotation events are logged and constitute input to the Scenario "
        "17 (Credential Compromise) recovery procedure documented in the Dashboard.",
    ]
    for item in cred_items:
        add_bullet(doc, item)

    add_heading(doc, "4.3.5  SeaweedFS Health Monitoring", level=3)
    add_para(doc,
        "A dedicated health check script (seaweedfs-k8s-filer-health.sh) is deployed "
        "in the platform operations toolkit. It validates the following on demand or "
        "from a scheduled Kubernetes CronJob:"
    )
    health_items = [
        "Filer pod readiness (liveness probe) and HTTP /healthz endpoint.",
        "Filer /status for metadata store health.",
        "Master /cluster/healthz for consensus health.",
        "Master /dir/status for writable volume slot headroom — triggers a warning "
        "when fewer than 2 free volume slots remain.",
        "df utilisation on each Volume Server pod — warning threshold at 85%, "
        "critical at 95%.",
        "df on the Filer pod for metadata store disk health.",
    ]
    for item in health_items:
        add_bullet(doc, item)
    add_para(doc,
        "Exit codes from the health check are: 0 (healthy), 1 (warning), "
        "2 (critical). These are consumed by PMM alerting and trigger the "
        "SeaweedFS-specific DR scenarios described in Section 8."
    )

    add_heading(doc, "4.4  Service Mesh and Network Security", level=2)
    add_para(doc,
        "All inter-service communication within and between clusters is protected by "
        "Istio mutual TLS (mTLS). In multi-cluster configurations, an East-West Gateway "
        "bridges cluster networks, enabling services in cluster-a (primary, network1) "
        "and cluster-b (DR, network2) to communicate securely across cluster boundaries "
        "without transmitting plain-text traffic."
    )
    add_para(doc,
        "The Istio control plane (istiod) manages certificate issuance and rotation "
        "for all service identities using short-lived SVIDs (~24-hour TTL). The DR "
        "cluster (cluster-b) is configured with Mesh ID mesh1, Cluster Name cluster-b, "
        "and Network network2. East-West Gateway ports: 15021 (health), 15443 (TLS), "
        "15012 (istiod discovery), 15017 (webhook)."
    )

    add_heading(doc, "4.5  Monitoring and Observability", level=2)
    add_para(doc,
        "The platform is monitored by Percona Monitoring and Management (PMM) deployed "
        "in the pmm namespace. PMM ingests metrics from Percona PMM Client agents on "
        "each PXC node and from Victoria Metrics Kubernetes stack (kube-state-metrics "
        "with custom resource state configuration). Each metric series carries a "
        "k8s_cluster_id label identifying the originating cluster."
    )
    add_para(doc, "Key alerting rules include:")
    alerts = [
        "MySQL instance down (critical).",
        "No MySQL instances monitored (indicates monitoring pipeline failure).",
        "PXC backup stale — no successful backup in 30 hours.",
        "SSL/TLS handshake failures at the application layer.",
        "SeaweedFS volume slot headroom below threshold (critical).",
        "SeaweedFS disk utilisation above 85% (warning) and 95% (critical).",
        "Replication lag on async DR replicas exceeding configured threshold.",
    ]
    for a in alerts:
        add_bullet(doc, a)
    add_page_break(doc)

    # ── SECTION 5: RECOVERY OBJECTIVES ──────────────────────────────────────
    add_heading(doc, "5  Recovery Objectives Framework", level=1)
    add_heading(doc, "5.1  Classification of Events", level=2)
    impact_tbl = doc.add_table(rows=1, cols=3)
    impact_tbl.style = "Table Grid"
    header_row(impact_tbl, "Impact Level", "Definition", "Examples")
    impact_rows = [
        ("Low",
         "Degraded or unavailable non-critical function; no data loss; no customer-visible outage.",
         "Single pod failure with immediate self-healing."),
        ("Medium",
         "Partial service degradation or brief customer-visible impact; limited data loss risk.",
         "Worker node failure; replication lag; tablespace exhaustion."),
        ("High",
         "Significant service impairment or risk of data loss; material customer impact.",
         "Quorum loss; credential compromise; DDL blocks writes; disk exhaustion."),
        ("Critical",
         "Complete service loss or confirmed data loss; potential regulatory notification required.",
         "Primary DC down; ransomware; accidental restore from wrong backup."),
    ]
    for i, row_data in enumerate(impact_rows):
        add_table_row(impact_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()

    add_heading(doc, "5.2  Recovery Objective Targets by Impact Level", level=2)
    rto_tbl = doc.add_table(rows=1, cols=4)
    rto_tbl.style = "Table Grid"
    header_row(rto_tbl, "Impact Level", "Target RTO", "Target RPO", "Backup Tier Required")
    rto_rows = [
        ("Low",      "≤ 10 minutes",  "0",           "In-cluster Galera replication"),
        ("Medium",   "≤ 60 minutes",  "≤ 5 minutes", "PITR-capable"),
        ("High",     "≤ 4 hours",     "≤ 15 minutes","Full backup + PITR"),
        ("Critical", "≤ 8 hours",     "≤ 120 seconds","Async DR replica or full backup"),
    ]
    for i, row_data in enumerate(rto_rows):
        add_table_row(rto_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()

    add_heading(doc, "5.3  Data Loss Tolerance", level=2)
    add_para(doc,
        "The platform's baseline commitment is zero data loss for any scenario in which "
        "the PXC cluster maintains quorum. Galera's synchronous certification protocol "
        "ensures that a transaction is not acknowledged to the client until it has been "
        "applied to a write-set on a quorum of nodes. Data loss scenarios are bounded as "
        "follows:"
    )
    loss_items = [
        ("Partial quorum events",
         "up to a small number of unflushed transactions in-flight at the moment of "
         "failure."),
        ("Async replica failover",
         "up to the replication lag at the moment of DC failure (typically ≤ 60 "
         "seconds under normal load)."),
        ("Logical data loss (DROP/DELETE/TRUNCATE)",
         "bounded by the PITR RPO (approximately 60–120 seconds in the worst case, "
         "given the 60-second binary log upload interval to SeaweedFS)."),
    ]
    for bold_part, rest in loss_items:
        add_bullet(doc, rest, bold_prefix=bold_part)
    add_page_break(doc)

    # ── SECTION 6: BACKUP STRATEGY ───────────────────────────────────────────
    add_heading(doc, "6  Backup and Data Protection Strategy", level=1)
    add_heading(doc, "6.1  Backup Architecture", level=2)
    add_para(doc,
        "All database backups are performed by Percona Backup for MySQL (PBM), "
        "orchestrated by the Percona Kubernetes Operator. Backups are written to the "
        "SeaweedFS Filer S3 endpoint. The backup target bucket is percona-backups "
        "within the SeaweedFS Filer. Backup credentials are managed as Kubernetes "
        "Secrets and are rotated through the Fleet GitOps pipeline."
    )

    add_heading(doc, "6.2  SeaweedFS Backup Integration", level=2)
    add_para(doc,
        "The Percona Operator backup storage configuration specifies the SeaweedFS "
        "Filer S3 endpoint as the storage target using standard S3-compatible "
        "parameters (endpointUrl, bucket, region, credentialsSecret). This design "
        "allows the Operator to treat SeaweedFS exactly as it would any S3-compatible "
        "object store, while all data remains on-premises."
    )
    add_para(doc,
        "The SeaweedFS Filer is the sole intermediary between the Operator backup jobs "
        "and the physical Volume Servers. The Filer translates incoming S3 PUT requests "
        "into needle writes distributed across the 3-Volume-Server pool according to "
        "the configured replication policy. The result is that each backup artefact has "
        "at least two physical copies on separate Volume Servers before the write is "
        "acknowledged."
    )
    add_para(doc,
        "The PITR tooling (pxc-restore / pitr-timestamp-finder) uses the SeaweedFS "
        "Filer HTTP API directly for binary log retrieval, enabling sub-minute "
        "timestamp resolution for accidental data-loss recovery."
    )

    add_heading(doc, "6.3  Backup Schedule and Retention", level=2)
    bkup_tbl = doc.add_table(rows=1, cols=5)
    bkup_tbl.style = "Table Grid"
    header_row(bkup_tbl, "Type", "Schedule (Cron)", "Retention", "Storage Target", "Encryption")
    bkup_rows = [
        ("Binlog (PITR)", "Continuous (60-sec upload)", "7 days",
         "SeaweedFS percona-backups bucket", "Yes"),
        ("Daily full (physical)", "0 2 * * * (02:00 daily)", "7 days",
         "SeaweedFS percona-backups bucket", "Yes"),
        ("Weekly full (physical)", "0 1 * * 0 (01:00 Sunday)", "8 weeks",
         "SeaweedFS percona-backups bucket", "Yes"),
        ("Monthly full (physical)", "30 1 1 * * (01:30 on 1st)", "12 months",
         "SeaweedFS percona-backups bucket", "Yes"),
    ]
    for i, row_data in enumerate(bkup_rows):
        add_table_row(bkup_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_para(doc,
        "Cross-site replication via SeaweedFS async replication controllers ensures "
        "that backup artefacts replicate to dr-std and dr-sec, making them available "
        "independently of primary-DC availability."
    )

    add_heading(doc, "6.4  Point-in-Time Recovery (PITR)", level=2)
    add_para(doc,
        "PITR is enabled on all production PXC clusters. Binary logs are uploaded to "
        "the SeaweedFS Filer at 60-second intervals. This provides an effective RPO of "
        "approximately 60–120 seconds in the worst case (one upload cycle missed plus "
        "detection latency)."
    )
    add_para(doc,
        "PITR restoration uses the pxc-restore tooling, which integrates with the "
        "SeaweedFS Filer HTTP API for binary log retrieval and provides:"
    )
    pitr_items = [
        "Interactive CLI listing all available backups with earliest and latest "
        "restorable timestamps.",
        "Automatic clone of cluster configuration to a target namespace (side-by-side "
        "restore without impacting the production cluster).",
        "Dry-run validation before execution.",
        "The pitr-timestamp-finder utility that scans binary logs on the SeaweedFS "
        "Filer to locate the exact transaction boundary immediately prior to a "
        "destructive event.",
    ]
    for item in pitr_items:
        add_bullet(doc, item)

    add_heading(doc, "6.5  Backup Validation", level=2)
    add_para(doc,
        "Silent backup failures — where jobs report success but artefacts are corrupt "
        "— are a latent risk. The following controls are in place:"
    )
    val_items = [
        ("Scheduled restore drills",
         "Quarterly automated and manual restore exercises. Results are tracked; any "
         "failure triggers an immediate P1 incident."),
        ("Checksum verification",
         "Backup manifests include checksums verified at restore time."),
        ("PMM alerting",
         "The PXC Backup Stale Critical alert fires when no successful backup has been "
         "recorded within a 30-hour window."),
        ("Backup count monitoring",
         "Automated checks verify the expected number of artefacts in the SeaweedFS "
         "bucket, catching premature deletion."),
    ]
    for bold_part, rest in val_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    add_heading(doc, "6.6  Data at Rest and in Transit", level=2)
    enc_items = [
        ("Encryption at rest",
         "Backup artefacts in SeaweedFS are encrypted at the Volume Server level. "
         "On-premises PVCs use encrypted storage volumes as configured by the "
         "infrastructure team."),
        ("Encryption in transit",
         "All traffic between PXC nodes uses SSL/TLS. All traffic between cluster "
         "services uses Istio mTLS. SeaweedFS API traffic (backup client to Filer, "
         "Filer to Volume Servers) is protected within the cluster network and "
         "optionally TLS-terminated."),
        ("Credential management",
         "All secrets (backup credentials, PMM tokens, database user credentials) are "
         "managed as Kubernetes Secrets, populated via the Fleet GitOps pipeline, and "
         "subject to the organisation's secret rotation policy."),
    ]
    for bold_part, rest in enc_items:
        add_bullet(doc, rest, bold_prefix=bold_part)
    add_page_break(doc)

    # ── SECTION 7: CERTIFICATES ──────────────────────────────────────────────
    add_heading(doc, "7  Certificate Management and Lifecycle", level=1)
    add_para(doc,
        "This section is addressed specifically to the CA/Browser Forum (CAB) "
        "Certificate Management Forum and to auditors requiring evidence of certificate "
        "lifecycle controls."
    )

    add_heading(doc, "7.1  Certificate Infrastructure Overview", level=2)
    add_para(doc, "The platform uses the following certificate authorities:", bold=True)
    add_para(doc, "Istio Service Mesh CA", bold=True)
    add_para(doc,
        "The Istio control plane (istiod) operates as an intermediate certificate "
        "authority for the service mesh. It issues short-lived X.509 SVIDs (~24-hour "
        "TTL) to every workload within the mesh, enabling mTLS between all services. "
        "These certificates are automatically rotated by Istio without service "
        "interruption. The root certificate and intermediate CA are generated at "
        "cluster bootstrap and stored in the cacerts Kubernetes Secret in the "
        "istio-system namespace."
    )
    ca_tbl = doc.add_table(rows=1, cols=2)
    ca_tbl.style = "Table Grid"
    header_row(ca_tbl, "File", "Purpose")
    for i, (f, p) in enumerate([
        ("ca-cert.pem", "Intermediate CA certificate (cluster-specific)"),
        ("root-cert.pem", "Shared root certificate across all clusters in the mesh"),
        ("cert-chain.pem", "Full certificate chain"),
        ("ca-key.pem", "CA private key (Kubernetes Secret only — never committed to Git)"),
    ]):
        add_table_row(ca_tbl, f, p, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_para(doc, "Kubernetes Ingress / Application Certificates", bold=True)
    add_para(doc,
        "TLS certificates for ingress endpoints (PMM, DR Dashboard, SeaweedFS external "
        "load balancer where applicable) are managed via Traefik's certificate "
        "provisioning on on-premises environments, or manually managed and renewed "
        "through the Fleet GitOps pipeline under CAB oversight."
    )

    add_heading(doc, "7.2  Certificate Inventory", level=2)
    add_para(doc,
        "The following table is the Certificate Inventory Register for CAB review. "
        "Full detail is in Appendix D."
    )
    cert_tbl = doc.add_table(rows=1, cols=6)
    cert_tbl.style = "Table Grid"
    header_row(cert_tbl, "Certificate", "Issuer", "Scope", "Rotation Method",
               "Review Owner", "CAB Required")
    cert_rows = [
        ("Istio Root CA", "Self-signed (bootstrap)", "All mesh clusters",
         "Manual (CAB approval required)", "Platform Engineering", "Yes"),
        ("Cluster-B Intermediate CA", "Istio Root CA", "cluster-b mesh identity",
         "Manual (CAB approval required)", "Platform Engineering", "Yes"),
        ("Istio workload SVIDs", "Cluster intermediate CA", "Per-service identity",
         "Automatic (istiod, ~24h TTL)", "Istio (automated)", "No"),
        ("Ingress TLS (on-prem)", "Traefik / self-signed", "Cluster ingress endpoints",
         "Manual or cert-manager", "Platform Engineering", "Yes"),
        ("MySQL TLS (PXC-internal)", "Percona Operator", "PXC node-to-node TLS",
         "Operator-managed", "Percona Operator", "No"),
        ("SeaweedFS Filer TLS (if enabled)", "Internal self-signed", "Filer S3 API endpoint",
         "Manual rotation via Fleet", "Platform Engineering", "Yes"),
    ]
    for i, row_data in enumerate(cert_rows):
        add_table_row(cert_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()

    add_heading(doc, "7.3  Certificate Rotation Procedures", level=2)
    add_para(doc, "Routine Rotation (automated)", bold=True)
    add_para(doc,
        "Istio workload SVIDs are rotated automatically by istiod with a default TTL "
        "of approximately 24 hours. No CAB approval is required."
    )
    add_para(doc, "Root CA Rotation (CAB-gated)", bold=True)
    root_steps = [
        "CAB Change Request submission at least 5 business days prior to rotation.",
        "Generate new root CA certificate using the generate-ca-certs script in a "
        "test environment; validate cross-cluster connectivity.",
        "Dual-root transition period: Istio supports a transition window during which "
        "both old and new root certificates are trusted simultaneously.",
        "Apply new cacerts secret to all clusters via Fleet GitOps commit, reviewed "
        "and approved in CAB.",
        "Monitor mTLS handshake errors in PMM/Kiali during the 24–48 hour transition.",
        "Remove old root from trust bundle once all workload certificates are re-issued.",
        "Update Certificate Inventory Register (Appendix D) and close the CAB change.",
    ]
    for i, step in enumerate(root_steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step).font.size = Pt(11)
    add_para(doc, "Emergency Certificate Rotation", bold=True)
    add_para(doc,
        "In the event of a suspected key compromise or CA failure, emergency rotation "
        "bypasses the standard 5-day lead time but still requires a documented CAB "
        "emergency change record within 24 hours of execution. The recovery procedure "
        "is documented in the Certificate Expiration / Revocation scenario accessible "
        "from the DR Dashboard at [DR_DASHBOARD_URL]."
    )

    add_heading(doc, "7.4  Certificate Expiry Monitoring", level=2)
    add_bullet(doc, "PMM alerting fires on SSL/TLS handshake failures detected at the "
               "application layer.")
    add_bullet(doc, "The DR Dashboard includes the certificate-expiration-or-revocation "
               "scenario with a 45-minute RTO.")
    add_bullet(doc, "All manually managed certificates must have expiry dates recorded "
               "in Appendix D with a minimum 30-day advance renewal reminder in the "
               "organisation's certificate monitoring system.")
    add_page_break(doc)

    # ── SECTION 8: SCENARIO REFERENCE ───────────────────────────────────────
    add_heading(doc, "8  Disaster Recovery Scenario Reference and Commentary", level=1)
    add_heading(doc, "8.1  Purpose of this Section", level=2)
    add_para(doc,
        "This section provides architectural and compliance commentary on each "
        "catalogued disaster recovery scenario for the on-premises data platform. "
        "It does not reproduce step-by-step recovery procedures. Those procedures "
        "are maintained exclusively in the DR Emergency Response Dashboard."
    )
    add_para(doc,
        "The authoritative source for all recovery runbooks is the DR Dashboard "
        "accessible at: [DR_DASHBOARD_URL]"
    )
    add_para(doc,
        "Each scenario in this section references the Dashboard by scenario name. "
        "Readers should use the Dashboard's scenario search or impact-sorted list "
        "to locate the corresponding runbook. This commentary explains why each "
        "scenario matters in the on-premises context, what architectural assumptions "
        "it depends on, and where it intersects with CFIUS or CAB compliance "
        "obligations."
    )

    add_heading(doc, "8.2  Scenario Matrix Summary", level=2)
    sc_tbl = doc.add_table(rows=1, cols=7)
    sc_tbl.style = "Table Grid"
    header_row(sc_tbl, "#", "Scenario", "Impact", "Likelihood", "RTO", "RPO",
               "Automated Test")
    scenarios_data = [
        (1, "Ransomware attack", "Critical", "Low", "8 hours", "120 sec", "No"),
        (2, "Primary data center is down", "Critical", "Low", "30 min", "60 sec", "No"),
        (3, "Accidental production restore from wrong backup", "Critical", "Low", "4 hours", "15 min", "No"),
        (4, "Schema change or DDL blocks writes", "High", "Medium", "30 min", "0", "Yes"),
        (5, "Cluster loses quorum (multiple PXC pods down)", "High", "Low", "90 min", "60 sec", "No"),
        (6, "Accidental DROP/DELETE/TRUNCATE", "High", "Medium", "4 hours", "5 min", "No"),
        (7, "Widespread data corruption (bad migration)", "High", "Low", "6 hours", "15 min", "No"),
        (8, "Application change causes data corruption", "High", "Medium", "8 hours", "24 hours", "No"),
        (9, "HAProxy endpoints inaccessible", "High", "Medium", "30 min", "0", "Yes"),
        (10, "Credential compromise (DB or SeaweedFS/S3 keys)", "High", "Medium", "120 min", "15 min", "No"),
        (11, "Certificate expiration or revocation", "High", "Medium", "45 min", "0", "No"),
        (12, "Database disk space exhaustion", "High", "Medium", "30 min", "0", "No"),
        (13, "Connection pool exhaustion (max_connections)", "High", "Medium", "15 min", "0", "No"),
        (14, "Memory exhaustion causing OOM kills", "High", "Medium", "20 min", "0", "No"),
        (15, "DNS resolution failure", "High", "Medium", "30 min", "0", "No"),
        (16, "Network policy misconfiguration blocking DB access", "High", "Medium", "30 min", "0", "No"),
        (17, "Increased API call volume / performance degradation", "High", "Medium", "60 min", "0", "No"),
        (18, "Application change causes performance degradation", "High", "Medium", "45 min", "0", "No"),
        (19, "Encryption key rotation failure", "High", "Low", "90 min", "0", "No"),
        (20, "Clock skew between nodes (replication issues)", "High", "Low", "60 min", "0", "No"),
        (21, "Application causing excessive replication lag", "Medium", "Medium", "4 hours", "0", "No"),
        (22, "Kubernetes worker node failure (VM host crash)", "Medium", "Medium", "20 min", "0", "Yes"),
        (23, "Percona Operator / CRD misconfiguration", "Medium", "Medium", "45 min", "0", "Yes"),
        (24, "Kubernetes control plane outage (API server down)", "Medium", "Low", "90 min", "0", "No"),
        (25, "Storage PVC corruption for a single PXC node", "Medium", "Low", "3 hours", "5 min", "No"),
        (26, "Primary DC network partition from secondary (WAN cut)", "Medium", "Medium", "0 (no failover)", "N/A", "No"),
        (27, "Both DCs up but replication stops (broken channel)", "Medium", "Medium", "60 min", "0", "No"),
        (28, "SeaweedFS backup target unavailable (regional / cred issue)", "Medium", "Medium", "0 (runtime OK)", "N/A", "No"),
        (29, "SeaweedFS service failure (backup target unavailable)", "Medium", "Medium", "0 (runtime OK)", "N/A", "No"),
        (30, "Monitoring and alerting system failure during incident", "Medium", "Medium", "N/A", "N/A", "No"),
        (31, "Temporary tablespace exhaustion", "Medium", "Medium", "15 min", "0", "No"),
        (32, "Single PXC or HAProxy pod failure", "Low", "Medium", "2 min", "0", "Yes"),
        (33, "Backups complete but non-restorable (silent failure)", "High", "Low", "4 hours", "15 min", "No"),
        (34, "Backup retention policy failure (premature deletion)", "Low", "Low", "4 hours", "15 min", "No"),
        (35, "Audit log corruption or loss (compliance violation)", "Low", "Low", "2 hours", "0", "No"),
    ]
    for i, row_data in enumerate(scenarios_data):
        add_table_row(sc_tbl, *[str(x) for x in row_data], shade=(i % 2 == 1))
    doc.add_paragraph()
    add_note(doc,
        "Step-by-step runbooks for every scenario above are available in the DR "
        "Emergency Response Dashboard at [DR_DASHBOARD_URL]. The commentary below "
        "provides on-premises architectural context and compliance observations only."
    )

    add_heading(doc, "8.3  Critical Impact Scenario Commentary", level=2)

    add_heading(doc, "Scenario 1 — Ransomware Attack", level=3)
    add_para(doc,
        "Architecture relevance: This is the highest-severity scenario for the "
        "on-premises environment precisely because all compute, storage, and network "
        "infrastructure is co-located within organisation-controlled data centres. A "
        "ransomware event that reaches VMware host storage can simultaneously encrypt "
        "Kubernetes node filesystems, PVC data, and SeaweedFS volume data. The "
        "existence of asynchronous SeaweedFS replication at dr-std and dr-sec is the "
        "primary control: if backups at the primary site are encrypted before the "
        "attack is detected, the replicated copies at DR sites may still be intact, "
        "provided network segmentation prevents lateral movement to the DR sites in "
        "time."
    )
    add_para(doc,
        "CFIUS relevance: The response to a ransomware event on national-security-"
        "relevant data may trigger CFIUS notification obligations. This scenario has "
        "the highest potential regulatory reporting surface of any in the catalogue. "
        "Recovery actions — including isolation of hosts, engagement of forensics, "
        "and data integrity verification — should be documented for regulatory review."
    )
    add_para(doc,
        "Recovery: See 'Ransomware attack' in the DR Dashboard at [DR_DASHBOARD_URL]."
    )

    add_heading(doc, "Scenario 2 — Primary Data Centre is Down", level=3)
    add_para(doc,
        "Architecture relevance: Full DC loss exercises every asynchronous component "
        "simultaneously. The PXC async replica at dr-std is the failover target. "
        "Simultaneously, the SeaweedFS replication controller at dr-std must be "
        "verified healthy before backup jobs are re-pointed to the dr-std Filer. "
        "If the SeaweedFS replica at dr-std has fallen behind at the moment of DC "
        "loss, any backups made in the replication gap must be reconstructed via "
        "PITR from the binlog stream."
    )
    add_para(doc,
        "CFIUS relevance: Geographic distribution of data to dr-std and dr-sec, "
        "both within organisation-controlled premises, is the primary data-sovereignty "
        "control that satisfies CFIUS geographic restriction requirements."
    )
    add_para(doc,
        "Recovery: See 'Primary data center is down' in the DR Dashboard."
    )

    add_heading(doc, "Scenario 3 — Accidental Production Restore from Wrong Backup", level=3)
    add_para(doc,
        "Architecture relevance: The pxc-restore tooling's side-namespace restore "
        "capability was specifically designed to mitigate this scenario. Restores "
        "should always be performed to a side namespace and validated before "
        "cutting over to production. The pitr-timestamp-finder scans SeaweedFS "
        "Filer binary log artefacts to identify the correct restore point."
    )
    add_para(doc,
        "Compliance relevance: Audit logs capturing PerconaXtraDBClusterRestore "
        "resource creation events are the primary evidence for reconstructing the "
        "incident timeline. Both CFIUS reviewers and internal auditors require this "
        "trail."
    )
    add_para(doc,
        "Recovery: See 'Accidental production restore from wrong backup or wrong point in time' "
        "in the DR Dashboard."
    )

    add_heading(doc, "8.4  High Impact Scenario Commentary", level=2)

    sc_high = [
        ("Scenario 4 — Schema Change or DDL Blocks Writes",
         "The PXC Galera replication model means that a metadata lock on one node "
         "propagates write-blocking across the whole cluster — unlike traditional "
         "single-primary MySQL where the write block is local. PMM slow-query "
         "alerting is the detection mechanism. This scenario has an automated "
         "resiliency test (test_dr_schema_change_or_ddl_blocks_writes.py).",
         "None specific."),
        ("Scenario 5 — Cluster Loses Quorum",
         "Quorum loss is the most technically sensitive recovery scenario in the "
         "Galera model. Selecting the node with the highest wsrep_last_committed "
         "value for bootstrap is critical — an incorrect choice creates a diverged "
         "cluster that is difficult to reconcile. The PodDisruptionBudget ensures "
         "that Kubernetes does not evict more than one PXC pod simultaneously "
         "during voluntary disruptions (node drains, upgrades).",
         "Any manual bootstrap procedure that bypasses the standard Galera "
         "certification protocol should be documented as a change record for "
         "CFIUS audit trail purposes."),
        ("Scenario 6 — Accidental DROP/DELETE/TRUNCATE",
         "This is the primary use case for the pitr-timestamp-finder tool, which "
         "scans SeaweedFS-stored binary logs to find the exact transaction boundary "
         "before the destructive statement. The 60-second binlog upload interval "
         "to SeaweedFS sets the effective lower bound on how recently PITR can "
         "recover.",
         "Data loss events may trigger regulatory notification obligations under "
         "applicable data protection law. The Data Protection Officer should be "
         "notified of any confirmed data loss immediately."),
        ("Scenario 7 — Widespread Data Corruption",
         "Pre-migration backup gates — mandatory successful backup verification "
         "before any schema-altering migration is permitted to run — are the primary "
         "preventive control. The SeaweedFS backup must be confirmed healthy via "
         "the backup count check and checksum verification before migration approval.",
         "None specific beyond data loss notification obligations."),
        ("Scenario 8 — Application Change Causes Data Corruption",
         "Gradual corruption scenarios are the most difficult to bound. Multi-week "
         "backup retention (12-month monthly backups) is the safety net. The "
         "SeaweedFS replication to dr-std and dr-sec ensures that even if the "
         "primary Filer is compromised by the corrupting application, clean copies "
         "of older backups are accessible from DR sites.",
         "Application-layer integrity checks that would flag corruption early are "
         "a recommended control gap to close."),
        ("Scenario 9 — HAProxy Endpoints Inaccessible",
         "HAProxy and ProxySQL are Kubernetes-native workloads. Failure of their "
         "Kubernetes Service Endpoints — not the pods themselves — is the most "
         "common root cause. Automated test: test_dr_ingressvip_failure.py.",
         "None specific."),
        ("Scenario 10 — Credential Compromise",
         "In the on-premises environment, SeaweedFS S3 credentials carry additional "
         "weight: compromised backup keys could allow an attacker to list, download, "
         "or delete backup artefacts. The credential rotation runbook covers both "
         "database user credentials and SeaweedFS S3 credentials. SIEM alerting on "
         "SeaweedFS access log anomalies is the detection mechanism.",
         "CFIUS requires documentation of all suspected credential exposure incidents, "
         "the response timeline, and any data exfiltration assessment."),
        ("Scenario 11 — Certificate Expiration or Revocation",
         "On-premises, certificates for the Istio mesh CA, Traefik ingress, and "
         "SeaweedFS Filer external load balancer are all manually managed (automated "
         "renewal is not available for all). The Certificate Inventory in Appendix D "
         "must be current; this is the primary evidence for CAB reviewers.",
         "Root CA and intermediate CA rotations are CAB-gated. Emergency rotations "
         "require a post-hoc emergency change record within 24 hours per the CAB "
         "Certificate Management policy."),
        ("Scenario 12 — Database Disk Space Exhaustion",
         "On-premises storage is finite and capacity management is the operator's "
         "responsibility (unlike elastic cloud volumes). PMM disk usage alerting "
         "and PVC expansion procedures must be current. SeaweedFS Volume Server "
         "disk exhaustion is a related but independent concern monitored by the "
         "SeaweedFS health check.",
         "None specific beyond operational hygiene."),
        ("Scenario 13 — Connection Pool Exhaustion",
         "ProxySQL multiplexes connections, providing a buffer between application "
         "connection storms and MySQL max_connections. Proper ProxySQL configuration "
         "is the primary preventive control.",
         "None specific."),
        ("Scenario 14 — Memory Exhaustion (OOM Kills)",
         "Kubernetes memory limits on PXC pods must be set correctly relative to "
         "InnoDB buffer pool size. OOM kills that restart PXC pods trigger IST "
         "re-sync from peers, which is expected and automatic. Disk-based swap is "
         "disabled for database workloads per Kubernetes and database best practice.",
         "None specific."),
        ("Scenario 15 — DNS Resolution Failure",
         "On-premises DNS for Kubernetes service discovery relies on CoreDNS. "
         "External DNS for VMware host addresses may involve a separate resolver. "
         "Failover to direct IP addresses is documented in the Dashboard runbook.",
         "None specific."),
        ("Scenario 16 — Network Policy Misconfiguration",
         "Kubernetes NetworkPolicy rules govern all traffic flows within the cluster. "
         "A misconfigured policy that blocks the backup client's access to the "
         "SeaweedFS Filer S3 endpoint will silently fail all backup jobs. Backup "
         "count monitoring will detect this within 30 hours.",
         "None specific."),
        ("Scenario 17 — Increased API Call Volume / Performance Degradation",
         "Scale-up in the on-premises environment means increasing the PXC CR size "
         "field via Fleet GitOps, which requests new VMs from VMware. This has longer "
         "lead time than cloud auto-scaling. Horizontal scaling is bounded by "
         "available VMware host capacity.",
         "None specific."),
        ("Scenario 18 — Application Change Causes Performance Degradation",
         "Application deployment rollback via the CD pipeline is the primary response. "
         "PMM slow-query analysis identifies the problematic query.",
         "None specific."),
        ("Scenario 19 — Encryption Key Rotation Failure",
         "Backup encryption keys are the highest-value keys in the on-premises "
         "environment. A failed key rotation that leaves backups unreadable eliminates "
         "the data-loss recovery capability until the rotation is corrected. Key "
         "backup is essential.",
         "Key management procedures should be documented for CFIUS review given the "
         "sensitivity of backup data."),
        ("Scenario 20 — Clock Skew Between Cluster Nodes",
         "VMware VM clock synchronisation relies on VMware Tools time sync or NTP. "
         "Galera's certification protocol uses monotonically increasing sequence "
         "numbers, not wall-clock time, so moderate clock skew does not immediately "
         "break replication — but binlog timestamps will be incorrect, affecting PITR "
         "accuracy. NTP monitoring is a required control.",
         "None specific."),
        ("Scenario 33 — Backups Complete but Non-Restorable",
         "This is a latent risk that only surfaces during a restore attempt. Quarterly "
         "restore drills against the SeaweedFS backup target are the only reliable "
         "detection method. PMM backup stale alerting detects job failures but cannot "
         "detect artefact corruption. A non-restorable backup discovered during an "
         "actual incident substantially increases RTO.",
         "Auditors will request evidence of restore drill execution and results. "
         "Completed drill reports should be retained for at least 12 months."),
    ]
    for title, arch_note, compliance_note in sc_high:
        add_heading(doc, title, level=3)
        add_para(doc, "Architecture note: " + arch_note)
        if compliance_note != "None specific.":
            add_para(doc, "Compliance note: " + compliance_note)
        add_para(doc, f"Recovery: See '{title.split(' — ', 1)[-1]}' in the DR Dashboard at [DR_DASHBOARD_URL].")

    add_heading(doc, "8.5  Medium Impact Scenario Commentary", level=2)

    sc_medium = [
        ("Scenario 22 — Kubernetes Worker Node Failure",
         "This scenario is the most frequently occurring in the catalogue. Kubernetes "
         "PodDisruptionBudgets and topology spread constraints ensure that PXC never "
         "loses quorum due to a single node eviction. The Percona Operator "
         "automatically re-joins the rescheduled PXC pod via IST. Automated test: "
         "test_dr_kubernetes_worker_node_failure.py."),
        ("Scenario 23 — Percona Operator / CRD Misconfiguration",
         "All CRD changes flow via Fleet GitOps and require pull-request review. "
         "Rollback is achieved by reverting the commit in Fleet. Automated test: "
         "test_dr_percona_operator_crd_misconfiguration.py."),
        ("Scenario 24 — Kubernetes Control Plane Outage",
         "Running pods continue to serve traffic when the Kubernetes API server is "
         "down. The data platform continues to operate normally. Risk is elevated "
         "for any operation requiring the Operator to reconcile (scale, backup, "
         "restore). etcd backups are the primary recovery enabler."),
        ("Scenario 25 — Storage PVC Corruption",
         "SeaweedFS is not involved in this scenario (database PVCs are separate "
         "from SeaweedFS volumes). Recovery is via IST or SST re-seeding from "
         "healthy Galera peers."),
        ("Scenario 26 — Primary DC Network Partition from Secondary",
         "The default policy is to remain primary in the current DC and queue "
         "async replication. SeaweedFS replication to DR sites is also queued "
         "during the partition and catches up when connectivity is restored."),
        ("Scenario 27 — Both DCs Up but Replication Stops",
         "Replication lag between primary and dr-std has a direct SeaweedFS "
         "analogue: if MySQL binlog replication stops, SeaweedFS backup artefacts "
         "at dr-std may still be current (they replicate independently). GTID "
         "resync is the primary repair procedure."),
        ("Scenario 28 — SeaweedFS Backup Target Unavailable",
         "The production cluster continues to serve live traffic when the SeaweedFS "
         "Filer is unavailable — database operations are not dependent on the backup "
         "path. The risk is to the backup pipeline: if the Filer is unavailable for "
         "longer than the binlog retention window (7 days), PITR capability degrades. "
         "Failover to the dr-std SeaweedFS Filer is the immediate response. "
         "Credential issues are the most common root cause and are handled by "
         "rotating the SeaweedFS S3 credentials in the Kubernetes Secret."),
        ("Scenario 29 — SeaweedFS Service Failure",
         "This is distinct from Scenario 28 in root cause: the SeaweedFS Filer pod "
         "or its underlying Volume Servers have failed, rather than a credential or "
         "network issue. Recovery involves restarting the SeaweedFS Filer deployment "
         "or statefulset and verifying that Volume Server connectivity is restored. "
         "The 3-master / 3-volume-server topology ensures that the loss of a single "
         "Volume Server does not interrupt reads or writes — data is replicated "
         "across at least two Volume Servers."),
        ("Scenario 30 — Monitoring and Alerting System Failure",
         "PMM failure does not affect the database. Manual diagnostic procedures "
         "via kubectl and direct MySQL queries are documented in the Dashboard "
         "runbook for use when PMM is unavailable."),
        ("Scenario 31 — Temporary Tablespace Exhaustion",
         "The key diagnostic indicator is a disk usage alert that fires and then "
         "auto-clears (temp files are cleaned when the query dies). A persistent "
         "alert indicates a long-running temp-table query. This is separate from "
         "SeaweedFS Volume Server disk exhaustion."),
        ("Scenario 21 — Application Causing Excessive Replication Lag",
         "Primary cluster operation is unaffected. The risk is to DR readiness: if "
         "the async replica at dr-std falls behind, the effective RPO for a "
         "DC-failover scenario degrades. Application query throttling and bulk "
         "operation scheduling are the primary controls."),
    ]
    for title, note in sc_medium:
        add_heading(doc, title, level=3)
        add_para(doc, note)
        add_para(doc, f"Recovery: See '{title.split(' — ', 1)[-1]}' in the DR Dashboard at [DR_DASHBOARD_URL].")

    add_heading(doc, "8.6  Low Impact Scenario Commentary", level=2)
    sc_low = [
        ("Scenario 32 — Single PXC or HAProxy Pod Failure",
         "Self-healing via Kubernetes and the Percona Operator. No manual "
         "intervention required unless the pod enters a permanent crash loop. "
         "Automated test: test_dr_single_mysql_pod_failure.py."),
        ("Scenario 34 — Backup Retention Policy Failure",
         "Premature deletion of SeaweedFS backup artefacts reduces the recovery "
         "window. Backup count monitoring and immutable volume configuration "
         "are the mitigating controls."),
        ("Scenario 35 — Audit Log Corruption or Loss",
         "This scenario directly impacts CFIUS and CAB compliance obligations. "
         "Audit log integrity checks and off-system shipping of audit logs are "
         "required controls. Any gap in the audit trail must be documented for "
         "compliance reviewers."),
    ]
    for title, note in sc_low:
        add_heading(doc, title, level=3)
        add_para(doc, note)
        add_para(doc, f"Recovery: See '{title.split(' — ', 1)[-1]}' in the DR Dashboard at [DR_DASHBOARD_URL].")

    add_page_break(doc)

    # ── SECTION 9: DASHBOARD ─────────────────────────────────────────────────
    add_heading(doc, "9  DR Emergency Response Dashboard", level=1)
    add_heading(doc, "9.1  Overview", level=2)
    add_para(doc,
        "The DR Emergency Response Dashboard is a purpose-built web application "
        "deployed to the on-premises Kubernetes cluster. It provides browser-based "
        "access to all recovery runbooks without requiring access to version control, "
        "wiki systems, or documentation portals that may be unavailable during an "
        "incident. The Dashboard is the authoritative and sole source of step-by-step "
        "recovery procedures for all scenarios described in Section 8."
    )
    add_para(doc, "The Dashboard is accessible at: [DR_DASHBOARD_URL]")
    add_para(doc, "Kubernetes deployment command:")
    add_para(doc, "    kubectl apply -f dr-dashboard/k8s/deployment-on-prem.yaml",
             size=10)

    add_heading(doc, "9.2  Architecture", level=2)
    arch_items = [
        ("Backend", "Go standard library only — no external dependencies. Startup time "
         "under 100ms. Memory usage approximately 10–20 MB."),
        ("Frontend", "Vanilla JavaScript and CSS3. No framework dependencies requiring "
         "CDN or network access during an incident."),
        ("Data source", "File-based — reads directly from the same disaster_scenarios.json "
         "files consumed by the automated testing framework. Single source of truth."),
        ("Deployment", "Kubernetes Deployment in the on-premises environment, independent "
         "of the database cluster."),
    ]
    for bold_part, rest in arch_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    add_heading(doc, "9.3  Features", level=2)
    feat_items = [
        "On-call contact information displayed at page load — no login required.",
        "Scenarios sorted by business impact (Critical → High → Medium → Low) and "
        "then by likelihood.",
        "Each scenario card expands to show: Overview (RTO, RPO, MTTR, likelihood, "
        "detection signals, affected components) and Recovery Process (full markdown "
        "runbook with copy-to-clipboard code blocks).",
        "Automatic incident detection script (detect-scenario.sh) analyses cluster "
        "state and suggests the most likely matching scenario.",
    ]
    for item in feat_items:
        add_bullet(doc, item)

    add_heading(doc, "9.4  Security Posture", level=2)
    sec_items = [
        "Read-only operations — the Dashboard cannot modify cluster state.",
        "Path traversal protection on all file-serving endpoints.",
        "No database — no SQL injection surface.",
        "Stateless — no session management or credential storage.",
        "Designed for internal network access only; authentication must be added "
        "before exposing beyond the cluster-internal network.",
    ]
    for item in sec_items:
        add_bullet(doc, item)

    add_heading(doc, "9.5  API Endpoints", level=2)
    api_tbl = doc.add_table(rows=1, cols=3)
    api_tbl.style = "Table Grid"
    header_row(api_tbl, "Endpoint", "Method", "Description")
    api_rows = [
        ("/", "GET", "Serves the dashboard UI"),
        ("/api/scenarios", "GET", "Returns full scenario catalogue as JSON"),
        ("/api/recovery-process?file={name}.md", "GET",
         "Returns markdown runbook content for the named scenario"),
        ("/static/*", "GET", "Static assets (JS, CSS)"),
    ]
    for i, row_data in enumerate(api_rows):
        add_table_row(api_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_page_break(doc)

    # ── SECTION 10: TESTING ──────────────────────────────────────────────────
    add_heading(doc, "10  Automated Testing and Validation", level=1)
    add_heading(doc, "10.1  Testing Philosophy", level=2)
    add_para(doc,
        "This plan does not rely solely on documentation for its assurances. Recovery "
        "procedures are validated through a three-tier testing regime:"
    )
    test_tiers = [
        ("Unit tests (~30 seconds; no cluster required)",
         "Configuration validation, YAML rendering, Helm template validation. Ensures "
         "infrastructure-as-code is syntactically and semantically correct."),
        ("Integration tests (~5–8 minutes; requires running cluster)",
         "Verify Kubernetes version compatibility, StorageClass availability, backup "
         "Secret existence, Operator status, anti-affinity rules, and resource "
         "request/limit correctness."),
        ("Resiliency tests (~30–60 minutes; with optional LitmusChaos)",
         "End-to-end validation of specific disaster scenarios, polling for cluster "
         "recovery within the scenario's MTTR target."),
    ]
    for bold_part, rest in test_tiers:
        add_bullet(doc, rest, bold_prefix=bold_part)

    add_heading(doc, "10.2  Automated Scenario Tests", level=2)
    auto_tbl = doc.add_table(rows=1, cols=4)
    auto_tbl.style = "Table Grid"
    header_row(auto_tbl, "Test File", "Scenario", "Chaos Type", "MTTR Target")
    auto_rows = [
        ("test_dr_single_mysql_pod_failure.py", "Single PXC pod failure",
         "pod-delete (pxc component)", "600 sec"),
        ("test_dr_kubernetes_worker_node_failure.py", "Worker node failure",
         "node-drain", "1,200 sec"),
        ("test_dr_percona_operator_crd_misconfiguration.py", "Operator pod failure",
         "pod-delete (operator deployment)", "900 sec"),
        ("test_dr_ingressvip_failure.py", "HAProxy / ProxySQL endpoint failure",
         "pod-delete (proxysql component)", "600 sec"),
        ("test_dr_schema_change_or_ddl_blocks_writes.py", "DDL blocking writes",
         "Controlled DDL lock injection", "1,800 sec"),
    ]
    for i, row_data in enumerate(auto_rows):
        add_table_row(auto_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()

    add_heading(doc, "10.3  Restore Drill Schedule", level=2)
    drill_tbl = doc.add_table(rows=1, cols=5)
    drill_tbl.style = "Table Grid"
    header_row(drill_tbl, "Drill", "Frequency", "Owner", "Last Completed", "Next Due")
    drill_rows = [
        ("Full backup restore to side namespace", "Quarterly",
         "Platform Engineering", "—", "—"),
        ("PITR restore to specific timestamp", "Quarterly",
         "Platform Engineering", "—", "—"),
        ("DR site failover drill (full DC simulation)", "Semi-annually",
         "Platform Engineering + Operations", "—", "—"),
        ("SeaweedFS Filer failover drill (primary to dr-std)", "Semi-annually",
         "Platform Engineering + Storage Ops", "—", "—"),
        ("Certificate rotation dry-run (test environment)",
         "Annually or before any root CA rotation",
         "Platform Engineering", "—", "—"),
    ]
    for i, row_data in enumerate(drill_rows):
        add_table_row(drill_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_note(doc,
        "Completed drill reports are retained for a minimum of 12 months and made "
        "available to auditors upon request."
    )
    add_page_break(doc)

    # ── SECTION 11: ROLES ────────────────────────────────────────────────────
    add_heading(doc, "11  Roles, Responsibilities, and Escalation", level=1)
    add_heading(doc, "11.1  Roles", level=2)
    roles_tbl = doc.add_table(rows=1, cols=2)
    roles_tbl.style = "Table Grid"
    header_row(roles_tbl, "Role", "Responsibilities")
    roles_rows = [
        ("On-Call Platform Engineer",
         "First responder for all database platform incidents. Uses DR Dashboard to "
         "identify scenario and execute runbook. Escalates within 15 minutes if "
         "scenario cannot be identified or containment actions have no effect."),
        ("Platform Engineering Lead",
         "Owns this DR plan (DR-PLN-002). Declares disaster events. Coordinates with "
         "downstream teams. Approves deviations from runbook procedures."),
        ("CISO",
         "Notified of any security-related scenario (credential compromise, ransomware, "
         "audit log loss). Authorises emergency certificate rotation."),
        ("Data Protection Officer",
         "Notified of any scenario with confirmed or potential data loss. Responsible "
         "for regulatory notification decisions under applicable data protection law."),
        ("Change Advisory Board (CAB)",
         "Reviews and approves root CA and intermediate CA rotations and other "
         "certificate management changes. Receives emergency change records for "
         "post-hoc review within 24 hours."),
        ("Database Administrator",
         "Subject matter expert for PXC-specific recovery (quorum bootstrap, PITR "
         "execution, replication repair)."),
        ("Storage Operations",
         "Subject matter expert for SeaweedFS recovery, including Filer failover, "
         "Volume Server repair, and cross-site replication restoration."),
    ]
    for i, row_data in enumerate(roles_rows):
        add_table_row(roles_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()

    add_heading(doc, "11.2  Escalation Path", level=2)
    esc_items = [
        "On-Call Engineer",
        "→ (15 min, no progress) Platform Engineering Lead",
        "→ (30 min, Critical or High impact) CISO (security) / DPO (data loss)",
        "→ (as appropriate) Vendor Support (Percona, VMware)",
    ]
    for item in esc_items:
        add_bullet(doc, item)

    add_heading(doc, "11.3  Communication During Incidents", level=2)
    add_bullet(doc, "All active incidents are tracked in the organisation's incident management system.")
    add_bullet(doc, "Status updates are communicated to affected application teams every 30 minutes.")
    add_bullet(doc, "Post-incident reviews (PIRs) are conducted within 5 business days for any Medium, High, or Critical incident.")
    add_page_break(doc)

    # ── SECTION 12: MAINTENANCE ──────────────────────────────────────────────
    add_heading(doc, "12  Plan Maintenance and Governance", level=1)
    add_heading(doc, "12.1  Review Cycle", level=2)
    add_para(doc, "This document is reviewed:")
    review_items = [
        "Annually as a scheduled review.",
        "Following any declared disaster event (within 30 days of resolution).",
        "Following any significant architectural change to the data platform.",
        "Before any CFIUS review submission or regulatory audit.",
        "When the SeaweedFS topology changes (master/volume/filer count, replication "
        "policy, new DR sites).",
    ]
    for item in review_items:
        add_bullet(doc, item)

    add_heading(doc, "12.2  Change Control", level=2)
    add_para(doc, "Changes to this document are subject to:")
    cc_items = [
        "Version control via Git commit in the platform repository.",
        "Review and sign-off by the Platform Engineering Lead.",
        "Notification to the CAB for any changes affecting Section 7 (Certificate Management).",
        "Distribution to all named roles in Section 11.",
    ]
    for item in cc_items:
        add_bullet(doc, item)

    add_heading(doc, "12.3  Scenario Catalogue Maintenance", level=2)
    add_para(doc,
        "The disaster scenario catalogue (disaster_scenarios.json) is the single source "
        "of truth for the DR Dashboard and for the scenario matrix in Section 8 of this "
        "document. Changes to the catalogue require: a pull request to the platform "
        "repository; a corresponding update to the associated recovery process markdown "
        "in dr-dashboard/recovery_processes/on-prem/; verification that the DR Dashboard "
        "correctly renders the updated scenario; and an update to Section 8 of this "
        "document on next review."
    )
    add_page_break(doc)

    # ── SECTION 13: COMPLIANCE ───────────────────────────────────────────────
    add_heading(doc, "13  Compliance Considerations", level=1)
    add_heading(doc, "13.1  CFIUS National Security Considerations", level=2)
    add_para(doc,
        "This section addresses the operational resilience controls relevant to CFIUS "
        "national security review of the data platform."
    )
    cfius_items = [
        ("Data Sovereignty and Geographic Controls",
         "All production data resides within [ORGANIZATION_NAME]-controlled "
         "infrastructure. The on-premises deployment operates exclusively within data "
         "centres under the direct physical and administrative control of the "
         "organisation. Cross-border or cross-region data transfers do not occur as "
         "part of normal backup or replication operations. SeaweedFS replication to "
         "dr-std and dr-sec moves data between organisation-controlled sites only."),
        ("Access Controls",
         "Access to all platform infrastructure requires authentication via the "
         "organisation's identity provider. Administrative access to database systems "
         "is restricted to named individuals with documented business need. All "
         "administrative actions are logged. SeaweedFS S3 credentials are managed as "
         "Kubernetes Secrets with least-privilege access, rotated via the GitOps "
         "pipeline."),
        ("Encryption",
         "All data at rest (database volumes, SeaweedFS backup artefacts) is "
         "encrypted. All data in transit between cluster components uses TLS. "
         "Service-to-service communication uses Istio mTLS — no plain-text traffic "
         "crosses network boundaries."),
        ("Audit Trail Integrity",
         "Database audit logging is enabled on all production PXC instances. Audit "
         "logs are shipped to centralised storage independent of the database system. "
         "Scenario 35 documents the recovery procedure for audit log corruption, "
         "including compliance notification obligations. Any audit trail gap must be "
         "disclosed to the CISO and DPO."),
        ("Operational Resilience",
         "The platform is designed with no single point of failure at the database "
         "tier. Synchronous Galera replication ensures no transaction is lost due to a "
         "single node failure. Asynchronous DR replicas at two geographically separated "
         "on-premises sites provide multi-site redundancy. SeaweedFS replication "
         "provides independent backup asset availability at each DR site."),
        ("Supply Chain and Software Provenance",
         "All platform software (Percona Operator, SeaweedFS Helm chart, Istio) is "
         "declared in version-pinned Nix flake and Helm chart configurations. "
         "Dependency versions are reproducible and auditable. No runtime dependencies "
         "on public package registries exist in the production environment."),
    ]
    for bold_part, rest in cfius_items:
        add_bullet(doc, rest, bold_prefix=bold_part)

    add_heading(doc, "13.2  CA/Browser Forum Certificate Management", level=2)
    add_para(doc,
        "This section summarises the certificate management posture for the CAB "
        "Certificate Management Forum."
    )
    cab_items = [
        ("Certificate Inventory",
         "The authoritative certificate inventory is maintained in Appendix D of this "
         "document and updated at each CAB forum."),
        ("Rotation Governance",
         "All root CA and intermediate CA rotations are CAB-gated changes requiring a "
         "formal change record with a minimum 5-business-day lead time, except in "
         "emergency scenarios (which require a post-hoc emergency change record within "
         "24 hours of execution)."),
        ("Automated vs. Manual Certificates",
         "Istio workload SVIDs (~24-hour TTL) are fully automated and do not require "
         "CAB oversight. All other certificates in the inventory are manually managed "
         "and subject to CAB governance."),
        ("DR Impact of Certificate Failures",
         "Scenario 11 (Certificate Expiration or Revocation) documents the specific "
         "recovery procedure for certificate-related outages, with a 45-minute RTO "
         "and step-by-step instructions accessible from the DR Dashboard."),
        ("Compliance with Certificate Policies",
         "All certificates issued by organisation-controlled CAs comply with the "
         "organisation's PKI policy. Third-party certificates comply with the issuing "
         "CA's Certificate Practice Statement."),
    ]
    for bold_part, rest in cab_items:
        add_bullet(doc, rest, bold_prefix=bold_part)
    add_page_break(doc)

    # ── APPENDIX A: CAPABILITIES ─────────────────────────────────────────────
    add_heading(doc, "Appendix A — On-Premises Platform Capabilities Inventory", level=1)
    add_para(doc,
        "Status indicators: Green (complete/operational), Yellow (in progress/partial), "
        "Red (not started/blocked)."
    )
    cap_tbl = doc.add_table(rows=1, cols=5)
    cap_tbl.style = "Table Grid"
    header_row(cap_tbl, "Capability", "Description", "Design", "Develop", "Operations")
    cap_rows = [
        ("On-prem S3-compatible storage",
         "SeaweedFS object storage exposing S3-compatible APIs for backups and blobs.",
         "Green", "Yellow", "Red"),
        ("Synchronous Percona PXC cluster",
         "PXC with Galera synchronous replication for strong consistency.",
         "Yellow", "Red", "Green"),
        ("HAProxy read-only endpoint (round-robin, all 4 clusters)",
         "Dedicated read load balancer across all cluster replicas.",
         "Red", "Green", "Yellow"),
        ("HAProxy write endpoint (no distribution, all 4 clusters)",
         "Single logical write entry point via HAProxy.",
         "Green", "Yellow", "Red"),
        ("dr-std PXC async replica (Active/Passive)",
         "Async PXC replica at standard DR site.",
         "Yellow", "Red", "Green"),
        ("dr-sec PXC async replica (Active/Passive)",
         "Secondary DR async PXC replica for geographic/regulatory separation.",
         "Red", "Green", "Yellow"),
        ("dr-std SeaweedFS async replication controller",
         "Controller orchestrating SeaweedFS replication at standard DR site.",
         "Green", "Yellow", "Red"),
        ("dr-sec SeaweedFS async replication controller",
         "Companion controller for secondary DR SeaweedFS replication.",
         "Yellow", "Red", "Green"),
        ("PXC backups S3 bucket controller",
         "Operator-managed workload coordinating backup streams into SeaweedFS S3.",
         "Red", "Green", "Yellow"),
        ("SeaweedFS Filer per cluster",
         "POSIX-compatible file API per cluster via SeaweedFS Filer.",
         "Green", "Yellow", "Red"),
    ]
    for i, row_data in enumerate(cap_rows):
        add_table_row(cap_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_note(doc,
        "Mixed status across capabilities reflects that the platform is under active "
        "development. Capabilities with Red Operations status represent areas requiring "
        "additional operational hardening before this DRP can be considered fully "
        "validated for those capabilities."
    )
    add_page_break(doc)

    # ── APPENDIX B: BACKUP MATRIX ─────────────────────────────────────────────
    add_heading(doc, "Appendix B — Backup Schedule and Retention Matrix", level=1)
    bm_tbl = doc.add_table(rows=1, cols=6)
    bm_tbl.style = "Table Grid"
    header_row(bm_tbl, "Backup Type", "Schedule", "Storage Target", "Retention",
               "Encryption", "Cross-DC Replication")
    bm_rows = [
        ("Binary log (PITR)", "Continuous (60-sec upload interval)",
         "SeaweedFS Filer / percona-backups", "7 days", "Yes",
         "Yes (SeaweedFS async replication)"),
        ("Daily full (physical)", "02:00 daily",
         "SeaweedFS Filer / percona-backups", "7 days", "Yes", "Yes"),
        ("Weekly full (physical)", "01:00 every Sunday",
         "SeaweedFS Filer / percona-backups", "8 weeks", "Yes", "Yes"),
        ("Monthly full (physical)", "01:30 on 1st of month",
         "SeaweedFS Filer / percona-backups", "12 months", "Yes", "Yes"),
    ]
    for i, row_data in enumerate(bm_rows):
        add_table_row(bm_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_para(doc,
        "PITR coverage: Continuous from the oldest available daily backup (7 days) to "
        "the most recent 60-second binlog checkpoint."
    )
    add_para(doc,
        "Backup credential: SeaweedFS S3 credentials stored in Kubernetes Secret in "
        "the percona namespace. Rotation schedule: quarterly or upon suspected "
        "compromise."
    )
    add_page_break(doc)

    # ── APPENDIX C: RISK REGISTER ────────────────────────────────────────────
    add_heading(doc, "Appendix C — Scenario Risk Register", level=1)
    rr_tbl = doc.add_table(rows=1, cols=6)
    rr_tbl.style = "Table Grid"
    header_row(rr_tbl, "#", "Scenario", "Likelihood", "Impact", "Key Mitigating Control",
               "Residual Risk")
    rr_rows = [
        (1, "Ransomware attack", "Low", "Critical",
         "SeaweedFS immutable volume writes; DR replication; EDR", "Medium"),
        (2, "Primary DC down", "Low", "Critical",
         "Async DR replica; SeaweedFS replication; switchover runbook", "Low-Medium"),
        (3, "Accidental production restore", "Low", "Critical",
         "Side-namespace restore; pitr-timestamp-finder; approval gates", "Low"),
        (4, "DDL blocks writes", "Medium", "High",
         "PMM slow-query alerting; kill procedures", "Medium"),
        (5, "Cluster loses quorum", "Low", "High",
         "3-node PXC; PodDisruptionBudgets; bootstrap runbook", "Low"),
        (6, "Accidental DROP/DELETE/TRUNCATE", "Medium", "High",
         "PITR on SeaweedFS; pitr-timestamp-finder; 7-day binlog retention", "Medium"),
        (7, "Widespread data corruption", "Low", "High",
         "Pre-migration backup gates; PITR; change windows", "Low"),
        (8, "Application data corruption", "Medium", "High",
         "Audit logging; integrity checks; PITR", "Medium"),
        (9, "HAProxy inaccessible", "Medium", "High",
         "Direct PXC bypass; automated test", "Low"),
        (10, "Credential compromise", "Medium", "High",
         "Secret rotation via Fleet; SIEM; MFA; SeaweedFS access log alerting", "Medium"),
        (11, "Certificate expiry/revocation", "Medium", "High",
         "PMM alerting; CAB governance; 45-min RTO runbook", "Low-Medium"),
        (12, "Disk space exhaustion", "Medium", "High",
         "PMM disk usage alerts; PVC expansion procedures", "Low"),
        (13, "Connection pool exhaustion", "Medium", "High",
         "PMM connection monitoring; max_connections tuning", "Low"),
        (14, "OOM kills", "Medium", "High",
         "Kubernetes memory limits; PMM OOM alerting", "Low"),
        (15, "DNS failure", "Medium", "High",
         "DNS monitoring; /etc/hosts fallback", "Low"),
        (16, "Network policy misconfiguration", "Medium", "High",
         "GitOps review; network policy backup", "Low"),
        (17, "Load-spike degradation", "Medium", "High",
         "GitOps-driven scaling (CR size field); PMM alerting", "Low"),
        (18, "App-change degradation", "Medium", "High",
         "Slow query alerting; app rollback", "Low"),
        (19, "Encryption key rotation failure", "Low", "High",
         "Key backup; rollback procedure", "Low"),
        (20, "Clock skew", "Low", "High",
         "NTP monitoring; VM clock synchronisation", "Low"),
        (21, "Excessive replication lag", "Medium", "Medium",
         "PMM replication lag alerting; query throttling", "Low"),
        (22, "Worker node failure", "Medium", "Medium",
         "Anti-affinity; PodDisruptionBudgets; automated test", "Low"),
        (23, "Operator misconfiguration", "Medium", "Medium",
         "GitOps rollback; automated test", "Low"),
        (24, "Control plane outage", "Low", "Medium",
         "etcd backups; Rancher re-provision", "Low"),
        (25, "PVC corruption (single node)", "Low", "Medium",
         "IST/SST re-seeding; 3-node redundancy", "Low"),
        (26, "WAN network partition", "Medium", "Medium",
         "No auto-failover policy; replication queuing", "Low"),
        (27, "Replication channel broken", "Medium", "Medium",
         "GTID resync; monitoring", "Low"),
        (28, "SeaweedFS backup target unavailable", "Medium", "Medium",
         "Secondary SeaweedFS Filer at DR site; local buffering", "Low"),
        (29, "SeaweedFS service failure", "Medium", "Medium",
         "3-master / 3-volume topology; pod restart; failover", "Low"),
        (30, "Monitoring system failure", "Medium", "Medium",
         "Manual kubectl procedures; backup monitoring", "Low"),
        (31, "Temp tablespace exhaustion", "Medium", "Medium",
         "Query kill procedures; tmpdir configuration", "Low"),
        (32, "Single pod failure", "Medium", "Low",
         "Auto-restart; Galera sync", "Very Low"),
        (33, "Non-restorable backups", "Low", "High",
         "Restore drills; checksum verification; PMM alerting", "Low"),
        (34, "Backup retention policy failure", "Low", "Low",
         "Retention monitoring; SeaweedFS immutable volumes", "Very Low"),
        (35, "Audit log corruption", "Low", "Low",
         "Separate backup; integrity checks; compliance notification", "Very Low"),
    ]
    for i, row_data in enumerate(rr_rows):
        add_table_row(rr_tbl, *[str(x) for x in row_data], shade=(i % 2 == 1))
    doc.add_paragraph()
    add_page_break(doc)

    # ── APPENDIX D: CERTIFICATE INVENTORY ───────────────────────────────────
    add_heading(doc, "Appendix D — Certificate Inventory and Lifecycle Register", level=1)
    add_para(doc,
        "This register is the authoritative certificate inventory for CAB Certificate "
        "Management Forum review. It must be updated whenever a certificate is issued, "
        "renewed, rotated, or revoked."
    )
    ci_tbl = doc.add_table(rows=1, cols=8)
    ci_tbl.style = "Table Grid"
    header_row(ci_tbl, "Certificate Name", "Type", "Issuer", "Issued",
               "Expiry", "Auto-Renew", "CAB Required", "Notes")
    ci_rows = [
        ("Istio Root CA", "Root CA", "Self-signed (bootstrap)",
         "—", "—", "No", "Yes",
         "Shared across all clusters in mesh1"),
        ("cluster-b Intermediate CA", "Intermediate CA", "Istio Root CA",
         "—", "—", "No", "Yes",
         "Stored in cacerts secret, istio-system ns"),
        ("Istio workload SVIDs", "Leaf / SVID", "cluster-b Intermediate CA",
         "Dynamic", "~24h", "Yes (istiod)", "No",
         "Auto-rotated; no manual intervention"),
        ("PMM Ingress TLS (on-prem)", "TLS Leaf", "Traefik (self-signed)",
         "—", "—", "No", "Yes",
         "Self-signed; update on renewal"),
        ("DR Dashboard TLS (on-prem)", "TLS Leaf", "Traefik (self-signed)",
         "—", "—", "No", "Yes",
         "Update on deployment"),
        ("PXC Node TLS (internal)", "TLS Leaf", "Percona Operator internal CA",
         "Dynamic", "Operator-managed", "Yes (Operator)", "No",
         "Auto-managed by Operator"),
        ("SeaweedFS Filer TLS (if enabled)", "TLS Leaf", "Internal self-signed",
         "—", "—", "No", "Yes",
         "Rotate with Filer upgrade; required if Filer S3 exposed externally"),
        ("SeaweedFS Master TLS (if enabled)", "TLS Leaf", "Internal self-signed",
         "—", "—", "No", "Yes",
         "Rotate with Master upgrade"),
    ]
    for i, row_data in enumerate(ci_rows):
        add_table_row(ci_tbl, *row_data, shade=(i % 2 == 1))
    doc.add_paragraph()
    add_note(doc,
        "Fields marked '—' must be populated by the Platform Engineering team upon "
        "first CAB forum review. Expiry dates for all manually managed certificates "
        "must be tracked with a minimum 30-day advance renewal reminder."
    )
    add_page_break(doc)

    # ── APPENDIX E: SEAWEEDFS HEALTH PROCEDURES ──────────────────────────────
    add_heading(doc, "Appendix E — SeaweedFS Health Check and Validation Procedures", level=1)
    add_para(doc,
        "This appendix documents the operational health check and validation procedures "
        "for the SeaweedFS distributed storage tier. These procedures are referenced "
        "by the SeaweedFS-specific DR scenarios and are run routinely as part of "
        "preventive maintenance."
    )

    add_heading(doc, "E.1  Health Check Script", level=2)
    add_para(doc,
        "The seaweedfs-k8s-filer-health.sh script performs comprehensive health "
        "validation. Run it as follows:"
    )
    add_para(doc,
        "    ./scripts/seaweedfs-k8s-filer-health.sh --namespace seaweedfs-primary",
        size=10
    )
    add_para(doc, "Exit codes: 0 = healthy, 1 = warning, 2 = critical.")

    add_heading(doc, "E.2  Manual Validation Checklist", level=2)
    check_items = [
        "Verify all 3 Master pods are Running and Ready.",
        "Verify all 3 Volume Server pods are Running and Ready.",
        "Verify the Filer pod is Running and Ready.",
        "Test Master /cluster/healthz endpoint from inside the cluster.",
        "Test Filer /healthz and /status endpoints.",
        "Check Master /dir/status for writable volume slot headroom (minimum 2 free).",
        "Run df on each Volume Server pod — no mount above 85% used.",
        "Test S3 API: aws s3 ls s3://percona-backups/ "
        "--endpoint-url http://seaweedfs-filer.[ns].svc.cluster.local:8333.",
        "Verify that the most recent Percona backup artefact is present in the bucket.",
        "Confirm replication controllers at dr-std and dr-sec show no error state.",
    ]
    for i, item in enumerate(check_items, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item).font.size = Pt(11)

    add_heading(doc, "E.3  Backup Artefact Verification", level=2)
    add_para(doc,
        "After every restore drill, record the following in the Restore Drill log:"
    )
    drill_checks = [
        "Backup source: SeaweedFS Filer namespace and endpoint used.",
        "Backup artefact name and timestamp selected for restore.",
        "Side-namespace restore completed successfully (yes/no).",
        "Data row count or integrity check result.",
        "PITR timestamp tested (if applicable).",
        "Total duration of restore operation.",
        "Any discrepancies or anomalies observed.",
    ]
    for item in drill_checks:
        add_bullet(doc, item)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    add_page_break(doc)
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end_p.add_run("End of Document")
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run(
        "Document ID: DR-PLN-002  |  Version 1.0  |  2026-06-17\n"
        "For questions, contact the Platform Engineering team."
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    return doc


if __name__ == "__main__":
    print("Building document...")
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
